#!/usr/bin/python3
from sys import argv
from collections import namedtuple
import re
import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
import matplotlib.pyplot as plt

SETTING_DATA_REGEX = '^#\+(.*?): (.*)$'
SETTING_NODATA_REGEX = '^#\+(.*)$'
HEADING_REGEX = '^(\*+) (.*)$'
IMAGE_REGEX = '^\[\[(.*?)\]\]$'
TABLE_REGEX = '^\| (.*) \|$'
TABLE_DIVIDER_REGEX = '^\|(\+|-)+\|$'
TODO_KEYWORDS = ('IDEA', 'NEXT', 'TODO', 'STARTED', 'DONE', 'CANCELLED')

Setting = namedtuple('Setting', ['name', 'data'])
Heading = namedtuple('Heading', ['level', 'text'])
Body = namedtuple('Body', ['text'])
BulletPoint = namedtuple('BulletPoint', ['text'])
Image = namedtuple('Image', ['path'])
Table = namedtuple('Table', ['data'])

def safe_index(d, key, default):
    try:
        return d[key]
    except KeyError:
        return default

def tokenize(lines):
    tokens = []
    body = ''
    table = []

    def end_body():
        nonlocal tokens, body
        body = body.rstrip(' \n')
        if body:
            tokens.append(Body(body))
            body = ''

    def end_table():
        nonlocal tokens, table
        if table != []:
            tokens.append(Table(table))
            table = []

    for l in lines:
        heading_match = re.findall(HEADING_REGEX, l)
        setting_data_match = re.findall(SETTING_DATA_REGEX, l)
        setting_nodata_match = re.findall(SETTING_NODATA_REGEX, l)
        image_match = re.findall(IMAGE_REGEX, l)
        table_match = re.findall(TABLE_REGEX, l)
        table_divider_match = re.findall(TABLE_DIVIDER_REGEX, l)

        if setting_data_match:
            end_body()
            end_table()

            setting = setting_data_match[0]
            setting_name = setting[0]

            if setting_name in ('OPTIONS', 'IMAGE_ATTR', 'PLOT'):
                setting_data_raw = setting[1]
                setting_data = {}
                # Parse individual key-value pairs
                for option in setting_data_raw.split(' '):
                    pair = option.split(':')
                    key = pair[0]
                    value = pair[1]
                    setting_data.update({key: value})
            else:
                setting_data = setting[1]

            tokens.append(Setting(setting_name, setting_data))

        elif setting_nodata_match:
            end_body()
            end_table()

            setting_name = setting_nodata_match[0]
            tokens.append(Setting(setting_name, None))

        elif heading_match:
            end_body()
            end_table()

            heading = heading_match[0]
            heading_level = len(heading[0])
            heading_text = heading[1]

            first_word = heading_text.split(' ')[0]
            if first_word in TODO_KEYWORDS:
                heading_text = heading_text.strip(first_word + ' ')

            tokens.append(Heading(heading_level, heading_text))

        elif image_match:
            end_body()
            end_table()

            image_path = image_match[0]
            tokens.append(Image(image_path))

        elif table_match:
            end_body()

            table_row = table_match[0]
            table_data = [r.strip() for r in table_row.split('|')]
            table.append(table_data)

        elif table_divider_match:
            # The first row of a table is assumed to be the header
            continue

        elif l.strip().startswith('-'):
            end_body()
            end_table()

            # Bullet point
            l = l.lstrip('- ')
            tokens.append(BulletPoint(l))
                
        else:
            # Remove leading whitespace
            l = l.strip()
            body += ('\n' if l == '' else l + ' ')

    body = body.rstrip(' \n')
    if body:
        tokens.append(Body(body))
        body = ''

    if table != []:
        tokens.append(Table(table))
        table = []
    return tokens

def edit_style(doc, style_name, font_name, font_size, font_color):
    style = doc.styles[style_name]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.color.rgb = RGBColor.from_string(font_color)

def add_style(doc, style_name, base_style, font_name, font_size, font_color):
    styles = doc.styles
    new_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    new_style.base_style = styles[base_style]
    font = new_style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.color.rgb = RGBColor.from_string(font_color)

def convert(tokens):
    doc = Document()
    edit_style(doc, 'Normal', 'Arial', 12, '000000')
    add_style(doc, 'H1', 'Heading 1', 'Arial', 18, '000000')
    add_style(doc, 'H2', 'Heading 2', 'Arial', 14, '000000')
    add_style(doc, 'H3', 'Heading 3', 'Arial', 12, '000000')

    line_spacing = 0
    image_scale = 0
    caption = ''
    plot = False
    plot_type = ''
    xlim = [0, 0]
    ylim = [0, 0]

    for t in tokens:
        if type(t) == Setting:
            if t.name == 'OPTIONS':
                line_spacing = float(safe_index(t.data, 'line-spacing', 2))
            elif t.name == 'IMAGE_ATTR':
                image_scale = float(safe_index(t.data, 'scale', 1))
            elif t.name == 'CAPTION':
                caption = t.data
            elif t.name == 'PAGE_BREAK':
                doc.add_page_break()
            elif t.name == 'PLOT':
                plot = True
                plot_type = t.data['type']
                xlim = safe_index(t.data, 'xlim', '0,0').split(',')
                xlim = [int(i) for i in xlim]
                ylim = safe_index(t.data, 'ylim', '0,0').split(',')
                ylim = [int(i) for i in ylim]

        elif type(t) == Heading:
            # Modifying the heading styles directly doesn't work;
            # use custom styles instead
            doc.add_paragraph(t.text, style='H'+str(t.level))

        elif type(t) == Body:
            doc.add_paragraph(t.text)
            para_format = doc.styles['Normal'].paragraph_format
            para_format.line_spacing = line_spacing

        elif type(t) == BulletPoint:
            para = doc.add_paragraph(t.text)
            para.style = 'List Bullet'

        elif type(t) == Image:
            doc.add_picture(t.path, width=Inches(5 * image_scale))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Caption
            para = doc.add_paragraph()
            para.add_run(caption).bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Blank line
            doc.add_paragraph()

        elif type(t) == Table:
            if plot:
                axes = plt.gca()
                axes.set_xlim(xlim)
                axes.set_ylim(ylim)

                plt.xlabel(t.data[0][0])
                plt.ylabel(t.data[0][1])

                # Convert data to form [[x1, ...], [y1, ...]]
                data = t.data[1:]
                data = list(map(list, zip(*data)))

                # Convert y values to ints
                for idx, i in enumerate(data[1]):
                    data[1][idx] = int(i)

                x_points = range(len(data[0]))
                if plot_type == 'line':
                    plt.plot(x_points, data[1])
                elif plot_type == 'bar':
                    plt.bar(x_points, data[1], align='center', width=0.3)

                plt.xticks(x_points, data[0])
                plt.tight_layout()
                plt.savefig('tmp.png')

                doc.add_picture('tmp.png')
                os.remove('tmp.png')

            else:
                table = doc.add_table(len(t.data), len(t.data[0]))
                table.style = 'TableGrid'

                for i, row in enumerate(t.data):
                    cells = table.rows[i].cells
                    for j, text in enumerate(row):
                        para = cells[j].paragraphs[0]
                        run = para.add_run(text)
                        if i == 0:
                            # Table headers
                            run.bold = True

            # Caption
            para = doc.add_paragraph()
            para.add_run(caption).bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Blank line
            doc.add_paragraph()
            plot = False

    doc.save(argv[2])

with open(argv[1]) as f:
    lines = f.read().split('\n')
    tokens = tokenize(lines)
    for t in tokens:
        print(t)
    convert(tokens)
